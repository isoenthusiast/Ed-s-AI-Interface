"""
📎 Attachment Handler — extracts text from uploaded files into markdown format.
Supports: PDF, DOCX, TXT, CSV
"""

import os
import csv
import io
import shutil
from pathlib import Path
from datetime import datetime

from src.config import ATTACHMENTS_DIR


# Ensure the attachments directory exists
ATTACHMENTS_DIR.mkdir(parents=True, exist_ok=True)


def save_attachment(filepath: str | Path) -> Path:
    """Copy a file into the attached/ folder with a timestamp prefix to avoid collisions.
    
    Returns the saved Path.
    """
    src = Path(filepath)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = f"{timestamp}_{src.name}"
    dest = ATTACHMENTS_DIR / safe_name
    shutil.copy2(src, dest)
    return dest


def extract_text(filepath: str | Path) -> str:
    """Extract text content from a file and return it as markdown.
    
    Args:
        filepath: Path to the file to extract.
    
    Returns:
        Markdown-formatted string of the file's content.
    """
    path = Path(filepath)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(path)
    elif suffix == ".docx":
        return _extract_docx(path)
    elif suffix in (".md", ".markdown"):
        return _extract_md(path)
    elif suffix == ".txt":
        return _extract_txt(path)
    elif suffix == ".csv":
        return _extract_csv(path)
    else:
        return f"*(Unsupported file type: {suffix})*"


def process_attachment(filepath: str | Path) -> dict:
    """Save and extract an attachment. Returns a dict with saved_path, filename, and markdown content."""
    path = Path(filepath)
    saved = save_attachment(path)
    content = extract_text(saved)
    return {
        "filename": path.name,
        "saved_path": str(saved),
        "content": content,
    }


# ═══════════════════════════════════════════════════════════════════
#  Extractors
# ═══════════════════════════════════════════════════════════════════

def _extract_pdf(path: Path) -> str:
    """Extract text from a PDF file."""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        return "*(PyPDF2 not installed. Run: pip install PyPDF2)*"

    try:
        reader = PdfReader(str(path))
        lines = [f"## 📄 {path.name}\n"]
        for i, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text and text.strip():
                lines.append(f"### Page {i}")
                lines.append(text.strip())
                lines.append("")
        if len(lines) == 1:
            return f"*(No extractable text found in {path.name})*"
        return "\n".join(lines)
    except Exception as e:
        return f"*(Failed to read PDF: {e})*"


def _extract_docx(path: Path) -> str:
    """Extract text from a DOCX file, preserving paragraph structure."""
    try:
        from docx import Document
    except ImportError:
        return "*(python-docx not installed. Run: pip install python-docx)*"

    try:
        doc = Document(str(path))
        lines = [f"## 📄 {path.name}\n"]

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                lines.append("")
                continue
            # Detect headings by style name
            style = para.style.name.lower() if para.style else ""
            if "heading 1" in style or "title" in style:
                lines.append(f"# {text}")
            elif "heading 2" in style:
                lines.append(f"## {text}")
            elif "heading 3" in style:
                lines.append(f"### {text}")
            elif "heading" in style:
                lines.append(f"**{text}**")
            else:
                lines.append(text)
            lines.append("")

        # Also extract tables
        for i, table in enumerate(doc.tables):
            lines.append(f"### Table {i + 1}")
            lines.append("")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
            if table.rows:
                # Add header separator after first row
                sep = "|" + "|".join(["---" for _ in table.rows[0].cells]) + "|"
                lines.insert(-len(table.rows), sep)
            lines.append("")

        if len(lines) == 1:
            return f"*(No content found in {path.name})*"
        return "\n".join(lines)
    except Exception as e:
        return f"*(Failed to read DOCX: {e})*"


def _extract_txt(path: Path) -> str:
    """Read a plain text file."""
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            text = path.read_text(encoding="latin-1")
        except Exception as e:
            return f"*(Failed to read TXT: {e})*"

    # Determine if it looks like code (heuristic)
    lines = text.split("\n")
    code_indicators = 0
    for line in lines[:50]:
        stripped = line.strip()
        if stripped and (stripped.startswith(("def ", "class ", "import ", "from ", "//", "/*", "#", "<?php", "<script", "function", "var ", "let ", "const ", "{")) or "=" in stripped):
            code_indicators += 1
    likely_code = code_indicators > 3 and len(lines) > 5

    if likely_code:
        # Try to guess language
        ext_map = {".py": "python", ".js": "javascript", ".ts": "typescript",
                   ".html": "html", ".css": "css", ".json": "json", ".xml": "xml",
                   ".sql": "sql", ".sh": "bash", ".bat": "batch", ".ps1": "powershell",
                   ".java": "java", ".cpp": "cpp", ".c": "c", ".cs": "csharp",
                   ".go": "go", ".rs": "rust", ".rb": "ruby", ".php": "php"}
        lang = ext_map.get(path.suffix.lower(), "")
        return f"## 📄 {path.name}\n\n```{lang}\n{text}\n```"
    else:
        return f"## 📄 {path.name}\n\n{text}"


def _extract_md(path: Path) -> str:
    """Read a Markdown file and return its content as-is."""
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            text = path.read_text(encoding="latin-1")
        except Exception as e:
            return f"*(Failed to read Markdown: {e})*"
    return f"## 📄 {path.name}\n\n{text}"


def _extract_csv(path: Path) -> str:
    """Read a CSV file and format as a markdown table."""
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            text = path.read_text(encoding="latin-1")
        except Exception as e:
            return f"*(Failed to read CSV: {e})*"

    try:
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)
    except Exception as e:
        return f"*(Failed to parse CSV: {e})*"

    if not rows:
        return f"## 📄 {path.name}\n\n*(Empty CSV file)*"

    lines = [f"## 📄 {path.name}\n"]

    # Limit to 100 rows to avoid huge messages
    max_rows = 100
    truncated = len(rows) > max_rows + 1  # +1 for header
    display_rows = rows[:max_rows + 1] if truncated else rows

    # Header row
    header = display_rows[0]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("|" + "|".join(["---" for _ in header]) + "|")

    # Data rows
    for row in display_rows[1:]:
        # Pad row to match header length
        padded = row + [""] * (len(header) - len(row))
        lines.append("| " + " | ".join(padded[:len(header)]) + " |")

    if truncated:
        lines.append(f"\n*(Showing {max_rows} of {len(rows) - 1} data rows)*")

    return "\n".join(lines)
